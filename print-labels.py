import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def read_sammlung(path_to_sammlung):
    sammlung = pd.read_excel(path_to_sammlung, sheet_name="Sammlung")
    return sammlung

def read_codes(path_to_codes):
    with open(path_to_codes, 'r') as file:
        lines = file.readlines()
        # Remove newline characters from each line
        codes = [line.strip() for line in lines]
    return codes

def add_label(doc, sammler, album, land, code, jv, jb, mv, mb, bv, bb):
    p = doc.add_paragraph()

    def add_bold_run(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True

    def add_regular_run(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = False

    # Add the text elements with formatting
    add_regular_run(p, "Sammler:  ")
    add_bold_run(p, f"{sammler}\t\t\t\t")
    add_regular_run(p, "Album:  ")
    add_bold_run(p, f"{album}\n")

    add_regular_run(p, "Sammelgebiet:\t")
    add_bold_run(p, f"{land}\n")

    add_regular_run(p, "Code:  ")
    add_bold_run(p, f"{code}\n")

    add_regular_run(p, "Zeitspanne:  ")
    add_bold_run(p, f"{jv} - {jb} \t")
    add_regular_run(p, "Mi-Nr.-Bereich:  ")
    add_bold_run(p, f"{mv} - {mb}\t")
    add_regular_run(p, "Blöcke:  ")
    add_bold_run(p, f"{bv} - {bb}\n")

    add_regular_run(p, "-------------------------------------------------------------------------------------------------------------")

    # Add a box around the text
    set_paragraph_border(p)

def set_paragraph_border(paragraph):
    p = paragraph._p  # Access the XML element for the paragraph
    pPr = p.get_or_add_pPr()
    # Set border properties
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 pt, adjust as needed
        border.set(qn('w:space'), '8')
        border.set(qn('w:color'), 'auto')
        pPr.append(border)

def make_codes(path_to_sammlung, path_to_codes, path_to_result):
    sammlung = read_sammlung(path_to_sammlung)
    codes = read_codes(path_to_codes)

    # get code information
    sammlung_filtered = sammlung[sammlung['Code'].isin(codes)].reset_index()

    # check if each code found its line in Sammlung.xls
    if sammlung_filtered.shape[0] != len(codes):
        raise Exception("Some codes could not be found or one code has two or more lines in Sammlung.xls")
    
    # create the result document
    doc = Document()

    # add labels
    for index, row in sammlung_filtered.iterrows():
        print(index)
        if index > 0 and index % 7 == 0:
            doc.add_page_break()
        add_label(doc, "BENINI Piero", row['Album'], row['Land'], row['Code'], row['Jv'], row['Jb'],
                  row['Mv'], row['Mb'], row['Bv'], row['Bb'])


    doc.save(path_to_result)


if __name__ == "__main__":

    ##########################
    ###### Variables #########
    ##########################

    path_to_sammlung = 'Sammlung.xls'

    path_to_codes = 'codes.txt'

    path_to_result = 'result.docx'

    ##########################
    ##########################
    ##########################

    make_codes(path_to_sammlung, path_to_codes, path_to_result)
